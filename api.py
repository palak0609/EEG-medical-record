from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.responses import JSONResponse, FileResponse
import os
import shutil
from typing import List
import tempfile
from datetime import datetime
import json
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import glob
import numpy as np
import zipfile

# Import the extraction functions from pdf_extraction.py
from pdf_extraction import extract_coherence_phase_lag, extract_other_sections, merge_data, save_to_csv

app = FastAPI(
    title="PDF Data Extraction API",
    description="API for extracting and processing data from PDF files",
    version="1.0.0"
)

# Create a directory to store uploaded files
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

@app.post("/extraction")
async def extract_data(
    pdf1: UploadFile = File(...),
    pdf2: UploadFile = File(...)
):
    """
    Upload 2 PDFs, extract data, perform all calculations, and return the processed results.
    """
    # Validate file types
    for pdf in [pdf1, pdf2]:
        if not pdf.filename.lower().endswith('.pdf'):
            raise HTTPException(
                status_code=400, 
                detail=f"File {pdf.filename} is not a PDF"
            )
    
    try:
        # Create a temporary directory for processing
        with tempfile.TemporaryDirectory() as temp_dir:
            pdf_paths = []
            
            # Save uploaded files
            for pdf in [pdf1, pdf2]:
                file_path = os.path.join(temp_dir, pdf.filename)
                with open(file_path, "wb") as buffer:
                    shutil.copyfileobj(pdf.file, buffer)
                pdf_paths.append(file_path)
            
            # Process the PDFs
            coherence_phase_lag_data = extract_coherence_phase_lag(pdf_paths)
            other_sections_data = extract_other_sections(pdf_paths)
            
            # Merge the data
            merged_data = merge_data(coherence_phase_lag_data + other_sections_data)
            
            # Create output directory with timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_dir = UPLOAD_DIR / timestamp
            output_dir.mkdir(exist_ok=True)
            
            # Save CSV files
            csv_output_path = output_dir / "extracted_data.csv"
            save_to_csv(merged_data, str(csv_output_path))
            
            # For each section CSV, apply calculations
            for section in merged_data.keys():
                section_csv_path = output_dir / f"{section}.csv"
                section_type = section
                apply_all_calculations(str(section_csv_path), section_type)
            
            # Create the combined document
            doc_output_path = output_dir / "eeg_analysis_summary.docx"
            create_combined_document(
                glob.glob(str(output_dir / "*.csv")),
                str(doc_output_path)
            )
            
            # Zip the output directory
            zip_path = str(output_dir) + ".zip"
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for root, dirs, files in os.walk(output_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arcname = os.path.relpath(file_path, start=output_dir)
                        zipf.write(file_path, arcname)
            
            # Return the zip file as a download
            return FileResponse(
                zip_path,
                filename=os.path.basename(zip_path),
                media_type='application/zip'
            )
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/")
async def root():
    """
    Root endpoint that returns API information
    """
    return {
        "message": "Welcome to PDF Data Extraction API",
        "endpoints": {
            "/extraction": "POST - Upload and process PDF files",
            "/docs": "GET - API documentation"
        }
    }

@app.post("/calculation")
async def calculation(
    csv_file: UploadFile = File(...),
    section_type: str = File(...)
):
    """
    Accepts a CSV file and section type, returns band-wise absolute sums.
    """
    # Save uploaded CSV to a temp file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".csv") as tmp:
        tmp.write(await csv_file.read())
        tmp_path = tmp.name

    try:
        summary = bandwise_absolute_sum(tmp_path, section_type)
        return {"summary": summary}
    finally:
        os.remove(tmp_path)

def create_combined_document(csv_files, output_path):
    """
    Create a single document combining all CSV files with metrics analysis
    """
    doc = Document()
    # Add custom heading at the top
    eeg_heading = doc.add_paragraph()
    run = eeg_heading.add_run("EEG//")
    run.bold = True
    run.font.size = Pt(18)
    # Default color (do not set run.font.color)
    doc.add_heading("Band-wise Metrics Analysis: ", level=1)

    # Define the desired order of sections
    section_order = [
        "Z Scored FFT Absolute Power",
        "Z Scored FFT Power Ratio",
        "Z Scored Peak Frequency",
        "Z Scored FFT Coherence",
        "Z Scored FFT Phase Lag"
    ]
    # Map section names to filenames
    section_to_file = {os.path.basename(f).replace('.csv', ''): f for f in csv_files}
    # Iterate in the desired order
    for section in section_order:
        file_path = section_to_file.get(section)
        if not file_path:
            continue
        section_name = section
        df = pd.read_csv(file_path)
        metrics_text = compute_subsection_bandwise_metrics(
            df,
            subsection_col="Subsection",
            band_col="Band", 
            set1_col="T1 Z",
            set2_col="T2 Z",
            normalize_col="Normalize"
        )
        section_heading = doc.add_heading(section_name, level=2)
        for block in metrics_text.strip().split("\n Subsection: "):
            if block.strip():
                if not block.startswith("Subsection:"):
                    block = "Subsection: " + block
                lines = block.strip().split("\n", 1)
                subsection_title = lines[0].replace("Subsection: ", "").strip()
                content = lines[1] if len(lines) > 1 else ""
                para = doc.add_paragraph()
                run = para.add_run(f"Subsection: {subsection_title}")
                run.bold = True
                doc.add_paragraph(content.strip(), style='Normal')
        doc.add_paragraph()
    doc.save(output_path)
    print(f"Combined analysis saved to {output_path}")

def process_all_csv_files(input_directory, output_doc="combined_analysis.docx"):
    """
    Process all CSV files in the input directory and create a combined document
    """
    # Get all CSV files in the directory
    csv_files = glob.glob(os.path.join(input_directory, "*.csv"))
    
    if not csv_files:
        print(f"No CSV files found in {input_directory}")
        return
    
    # Create the combined document
    create_combined_document(csv_files, output_doc)

def main():
    # Example usage
    input_directory = "uploads"  # Directory containing your CSV files
    output_doc = "eeg_analysis_summary.docx"
    
    process_all_csv_files(input_directory, output_doc)

def bandwise_absolute_sum(csv_path: str, section_type: str) -> dict:
    """Calculate band-wise absolute sums for a given CSV file"""
    df = pd.read_csv(csv_path)
    summary = {
        'section_type': section_type,
        'subsections': {}
    }
    
    for subsection in df['Subsection'].unique():
        sub_df = df[df['Subsection'] == subsection]
        summary['subsections'][subsection] = {}
        
        for band in sub_df['Band'].unique():
            band_df = sub_df[sub_df['Band'] == band]
            t1_sum = band_df['T1 Z'].abs().sum()
            t2_sum = band_df['T2 Z'].abs().sum()
            summary['subsections'][subsection][band] = {
                'T1_sum': t1_sum,
                'T2_sum': t2_sum
            }
    
    return summary

def apply_all_calculations(csv_path: str, section_type: str) -> None:
    """Apply all calculations to a CSV file and save results"""
    df = pd.read_csv(csv_path)
    
    # Add section type as a column
    df['Section_Type'] = section_type
    
    # Calculate metrics
    for subsection in df['Subsection'].unique():
        sub_df = df[df['Subsection'] == subsection]
        for band in sub_df['Band'].unique():
            band_df = sub_df[sub_df['Band'] == band]
            
            # Calculate absolute sums
            t1_sum = band_df['T1 Z'].abs().sum()
            t2_sum = band_df['T2 Z'].abs().sum()
            
            # Calculate averages
            t1_avg = band_df['T1 Z'].abs().mean()
            t2_avg = band_df['T2 Z'].abs().mean()
            
            # Calculate delta and percent change
            delta = abs(t1_avg - t2_avg)
            percent_change = (delta / t1_avg) * 100 if t1_avg != 0 else 0
            
            # Update the DataFrame with calculations
            mask = (df['Subsection'] == subsection) & (df['Band'] == band)
            df.loc[mask, 'T1_Sum'] = t1_sum
            df.loc[mask, 'T2_Sum'] = t2_sum
            df.loc[mask, 'Delta'] = delta
            df.loc[mask, 'Percent_Change'] = percent_change
    
    # Save the updated DataFrame
    df.to_csv(csv_path, index=False)

def compute_subsection_bandwise_metrics(df: pd.DataFrame, subsection_col: str, band_col: str, set1_col: str, set2_col: str, normalize_col: str) -> str:
    """Compute metrics for each subsection and band"""
    subsections = df[subsection_col].dropna().unique()
    result_blocks = []

    df[set1_col] = pd.to_numeric(df[set1_col], errors='coerce')
    df[set2_col] = pd.to_numeric(df[set2_col], errors='coerce')

    for subsection in subsections:
        subsection_df = df[df[subsection_col] == subsection]
        bands = subsection_df[band_col].dropna().unique()

        result_blocks.append(f"\n Subsection: {subsection} \n")

        for band in bands:
            band_df = subsection_df[subsection_df[band_col] == band][[set1_col, set2_col, normalize_col]].dropna()

            if band_df.empty:
                continue

            abs_sum_1 = np.abs(band_df[set1_col]).sum()
            abs_avg_1 = np.abs(band_df[set1_col]).mean()
            abs_sum_2 = np.abs(band_df[set2_col]).sum()
            abs_avg_2 = np.abs(band_df[set2_col]).mean()

            delta = abs(abs_avg_1 - abs_avg_2)
            percent_change = (delta / abs_avg_1) * 100 if abs_avg_1 != 0 else 0
            direction = "increase" if abs_avg_2 > abs_avg_1 else "decrease"

            total_rows = len(band_df)
            normalize_counts = band_df[normalize_col].astype(str).value_counts()
            normalize_yes = normalize_counts.get("Yes", 0)
            normalize_no = normalize_counts.get("No", 0)
            normalize_ns = normalize_counts.get("NS", 0)

            block = f""" Band: {band}
Set 1 ({set1_col}):
  Absolute Sum: {abs_sum_1:.2f}
  Average Absolute Value: {abs_avg_1:.3f}

Set 2 ({set2_col}):
  Absolute Sum: {abs_sum_2:.2f}
  Average Absolute Value: {abs_avg_2:.3f}

Differences:
  Delta: {delta:.3f}
  Percent Change: {percent_change:.2f}% ({direction} from Set 1 to Set 2)

Normalize Counts:
  Total Rows: {total_rows}
  "Normalize = Yes": {normalize_yes} ({(normalize_yes / total_rows) * 100 if total_rows > 0 else 0:.2f}%)
  "Normalize = No": {normalize_no} ({(normalize_no / total_rows) * 100 if total_rows > 0 else 0:.2f}%)
  "Normalize = NS": {normalize_ns} ({(normalize_ns / total_rows) * 100 if total_rows > 0 else 0:.2f}%)

__________________________________________________\n"""

            result_blocks.append(block)

    return "\n".join(result_blocks)

@app.get("/download/{timestamp}/{filename}")
async def download_file(timestamp: str, filename: str):
    file_path = UPLOAD_DIR / timestamp / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(str(file_path), filename=filename)

if __name__ == "__main__":
    main() 