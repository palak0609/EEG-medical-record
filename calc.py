from langchain_groq import ChatGroq
groq_api_key= "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"

llm = ChatGroq(groq_api_key=groq_api_key, model_name='llama-3.3-70b-versatile')

import pandas as pd
import numpy as np
from docx import Document
import os
from docx.shared import Pt

def save_to_docx(text: str, file_name: str):
    doc = Document()
    # Add custom heading at the top
    eeg_heading = doc.add_paragraph()
    run = eeg_heading.add_run("EEG//")
    run.bold = True
    run.font.size = Pt(18)
    # Default color (do not set run.font.color)
    doc.add_heading("Band-wise Metrics Analysis: ", level=1)

    # Metrics to bullet
    bullet_metrics = [
        "Absolute Sum:",
        "Average Absolute Value:",
        "Delta:",
        "Percent Change:",
        "Total Rows:",
        '"Normalize = Yes":',
        '"Normalize = No":',
        '"Normalize = NS":',
    ]

    for block in text.strip().split("\n Subsection: "):
        if block.strip():
            if not block.startswith("Subsection:"):
                block = "Subsection: " + block
            lines = block.strip().split("\n", 1)
            subsection_title = lines[0].replace("Subsection: ", "").strip()
            content = lines[1] if len(lines) > 1 else ""

            para = doc.add_paragraph()
            run = para.add_run(f"Subsection: {subsection_title}")
            run.bold = True

            # Add each line, prepending a bullet if it's a metric
            for line in content.strip().split("\n"):
                line_strip = line.strip()
                if any(line_strip.lstrip().startswith(metric) for metric in bullet_metrics):
                    doc.add_paragraph(f"• {line_strip}", style='Normal')
                elif line_strip:
                    doc.add_paragraph(line_strip, style='Normal')

    doc.save(file_name)

def load_csv(file_path: str) -> pd.DataFrame:
    return pd.read_csv(file_path)

def compute_subsection_bandwise_metrics(df: pd.DataFrame, subsection_col: str, band_col: str, set1_col: str, set2_col: str, normalize_col: str) -> str:
    subsections = df[subsection_col].dropna().unique()
    result_blocks = []

    df[set1_col] = pd.to_numeric(df[set1_col], errors='coerce')
    df[set2_col] = pd.to_numeric(df[set2_col], errors='coerce')

    for subsection in subsections:
        subsection_df = df[df[subsection_col] == subsection]
        bands = subsection_df[band_col].dropna().unique()

        result_blocks.append(f"\n Subsection: {subsection} \n")

        for band in bands:
            band_df = subsection_df[subsection_df[band_col] == band][[set1_col, set2_col, normalize_col]].dropna()

            if band_df.empty:
                continue

            abs_sum_1 = np.abs(band_df[set1_col]).sum()
            abs_avg_1 = np.abs(band_df[set1_col]).mean()
            abs_sum_2 = np.abs(band_df[set2_col]).sum()
            abs_avg_2 = np.abs(band_df[set2_col]).mean()

            delta = abs(abs_avg_1 - abs_avg_2)
            percent_change = (delta / abs_avg_1) * 100 if abs_avg_1 != 0 else 0
            direction = "increase" if abs_avg_2 > abs_avg_1 else "decrease"

            total_rows = len(band_df)
            normalize_counts = band_df[normalize_col].astype(str).value_counts()
            normalize_yes = normalize_counts.get("Yes", 0)
            normalize_no = normalize_counts.get("No", 0)
            normalize_ns = normalize_counts.get("NS", 0)

            block = f""" Band: {band}
Set 1 ({set1_col}):
  Absolute Sum: {abs_sum_1:.2f}
  Average Absolute Value: {abs_avg_1:.3f}

Set 2 ({set2_col}):
  Absolute Sum: {abs_sum_2:.2f}
  Average Absolute Value: {abs_avg_2:.3f}

Differences:
  Delta: {delta:.3f}
  Percent Change: {percent_change:.2f}% ({direction} from Set 1 to Set 2)

Normalize Counts:
  Total Rows: {total_rows}
  "Normalize = Yes": {normalize_yes} ({(normalize_yes / total_rows) * 100 if total_rows > 0 else 0:.2f}%)
  "Normalize = No": {normalize_no} ({(normalize_no / total_rows) * 100 if total_rows > 0 else 0:.2f}%)
  "Normalize = NS": {normalize_ns} ({(normalize_ns / total_rows) * 100 if total_rows > 0 else 0:.2f}%)

__________________________________________________\n"""

            result_blocks.append(block)

    return "\n".join(result_blocks)

# Main execution
if __name__ == "__main__":
    import glob

    folder_path = "/content/drive/MyDrive/EEG/set7"  # <- change this to your folder
    subsection_col = "Subsection"
    output_path = "/content/drive/MyDrive/EEG/res_7"
    band_col = "Band"
    set1_col = "T1 Z"
    set2_col = "T2 Z"
    normalize_col = "Normalize"
    if not os.path.exists(output_path):
        os.makedirs(output_path)

    excel_files = glob.glob(os.path.join(folder_path, "*.csv"))

    for file_path in excel_files:
        print(f"Processing: {os.path.basename(file_path)}")

        df = load_csv(file_path)
        metrics_text = compute_subsection_bandwise_metrics(df, subsection_col, band_col, set1_col, set2_col, normalize_col)

        output_file_name = os.path.splitext(os.path.basename(file_path))[0] + ".docx"
        save_to_docx(metrics_text, os.path.join(output_path, output_file_name))

        print(f"Saved: {output_file_name}")
